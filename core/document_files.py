import io
import re
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# reportlab по умолчанию не поддерживает кириллицу (base14-шрифты вроде Helvetica
# не содержат нужных глифов) — используем DejaVu Sans (assets/fonts/, см. README.md
# там же про лицензию).
_FONTS_DIR = Path(__file__).resolve().parent.parent / "assets" / "fonts"
_PDF_FONT_REGISTERED = False

_INVALID_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]')


def _register_pdf_fonts() -> None:
    global _PDF_FONT_REGISTERED
    if _PDF_FONT_REGISTERED:
        return
    pdfmetrics.registerFont(TTFont("DejaVuSans", str(_FONTS_DIR / "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(_FONTS_DIR / "DejaVuSans-Bold.ttf")))
    _PDF_FONT_REGISTERED = True


def _safe_filename(title: str) -> str:
    cleaned = _INVALID_FILENAME_CHARS.sub("", title).strip()
    return cleaned[:60] or "Документ"


def build_docx(title: str, body: str) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=1)
    for paragraph in body.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_xlsx(body: str) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    for line in body.strip().splitlines():
        row = [cell.strip() for cell in line.split("|")]
        sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def build_pdf(title: str, body: str) -> bytes:
    _register_pdf_fonts()
    buffer = io.BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=A4)
    title_style = ParagraphStyle("title", fontName="DejaVuSans-Bold", fontSize=16, spaceAfter=12)
    body_style = ParagraphStyle("body", fontName="DejaVuSans", fontSize=11, spaceAfter=10, leading=15)

    elements = [Paragraph(xml_escape(title), title_style), Spacer(1, 12)]
    for paragraph in body.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            escaped = xml_escape(paragraph).replace("\n", "<br/>")
            elements.append(Paragraph(escaped, body_style))
            elements.append(Spacer(1, 8))

    document.build(elements)
    return buffer.getvalue()


def build_document_file(format_: str, title: str, content: str) -> tuple[bytes, str]:
    """Собирает документ в реальный файл. Возвращает (байты файла, имя файла с
    расширением).

    Для docx/pdf content — текст с абзацами, разделёнными "\\n\\n". Для xlsx —
    таблица: одна строка текста на строку таблицы, ячейки разделены "|" (простой
    формат, который надёжно генерирует LLM — без вложенного JSON для таблиц).
    """
    filename = _safe_filename(title)

    if format_ == "docx":
        return build_docx(title, content), f"{filename}.docx"
    if format_ == "xlsx":
        return build_xlsx(content), f"{filename}.xlsx"
    if format_ == "pdf":
        return build_pdf(title, content), f"{filename}.pdf"

    raise ValueError(f"Неизвестный формат документа: {format_!r}")
